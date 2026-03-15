"""
サンプル Word テンプレート生成スクリプト
このスクリプトを一度だけ実行すると「template_sample.docx」が生成されます。
生成されたファイルを Word で開いて、自社の書式に合わせて編集してください。

実行方法:
  python create_template.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_sample_template(output_path: str = "template_sample.docx"):
    doc = Document()

    # ページ余白を設定
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # タイトル
    title = doc.add_heading("旅行申込書", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    doc.add_paragraph()

    # 作成日
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_date.add_run("書類作成日：{{ today }}")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # 確認番号
    p_conf = doc.add_paragraph()
    run = p_conf.add_run("予約確認番号：{{ confirmation_code }}")
    run.font.size = Pt(11)
    run.font.bold = True

    doc.add_paragraph()

    # 情報テーブル
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    rows_data = [
        ("氏名（漢字）", "{{ full_name }}"),
        ("住所", "{{ address }}"),
        ("電話番号", "{{ phone }}"),
        ("生年月日", "{{ date_of_birth }}"),
        ("メールアドレス", "{{ email }}"),
        ("予約日", "{{ booking_date }}"),
    ]

    for i, (label, placeholder) in enumerate(rows_data):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]

        # 幅の設定
        label_cell.width = Cm(5)
        value_cell.width = Cm(11)

        # ラベル（太字・背景色）
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.bold = True
        label_run.font.size = Pt(10.5)
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # プレースホルダー
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(placeholder)
        value_run.font.size = Pt(10.5)

    doc.add_paragraph()
    doc.add_paragraph()

    # 注意書き
    note = doc.add_paragraph()
    note_run = note.add_run(
        "※ この書類は自動生成されたものです。内容をご確認の上、署名・捺印してご提出ください。"
    )
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.save(output_path)
    print(f"サンプルテンプレートを作成しました: {output_path}")
    print()
    print("次のステップ:")
    print("  1. template_sample.docx を Word で開く")
    print("  2. 自社のロゴや書式に合わせてデザインを編集する")
    print("  3. {{ full_name }} などのプレースホルダーはそのまま残す")
    print("  4. 編集後のファイルを app.py から読み込むテンプレートとして使用する")


if __name__ == "__main__":
    create_sample_template()
